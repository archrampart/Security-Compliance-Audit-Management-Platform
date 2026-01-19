"""
Word (.docx) report generator for audit reports
Creates professional audit reports with cover page following international standards
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from sqlalchemy.orm import Session
from app.models.audit import Audit
from app.models.template import Severity, Status
from datetime import datetime
import tempfile
import os
from collections import Counter
from app.core.config import settings

# Image extensions that can be embedded in Word
IMAGE_EXTENSIONS = {'.jpg', '.jpeg', '.png', '.gif', '.webp'}

# i18n Translations for reports
TRANSLATIONS = {
    "tr": {
        "audit_report": "DENETİM RAPORU",
        "executive_summary": "ÖZET YÖNETİCİ RAPORU",
        "table_of_contents": "İÇİNDEKİLER",
        "scope_methodology": "DENETİM KAPSAMI VE METODOLOJİSİ",
        "findings": "BULGULAR",
        "conclusion": "SONUÇ VE ÖNERİLER",
        "appendices": "EKLER",
        "standard": "Standart",
        "organization": "Organizasyon",
        "project": "Proje",
        "audit_date": "Denetim Tarihi",
        "report_date": "Rapor Tarihi",
        "not_specified": "Belirtilmemiş",
        "finding_distribution": "Bulgu Dağılımı:",
        "finding": "bulgu",
        "no_findings": "Denetim kapsamında bulgu tespit edilmemiştir.",
        "finding_title": "BULGU",
        "severity": "Önem Derecesi",
        "status_label": "Durum",
        "control_reference": "Kontrol Referansı",
        "description": "Açıklama",
        "recommendation": "Öneri",
        "evidence_count": "Kanıt Sayısı",
        "evidence_desc": "Açıklama",
        "scope": "Denetim Kapsamı",
        "scope_text": "Bu denetim {standard} standartı kapsamında gerçekleştirilmiştir.",
        "summary_text": "Bu rapor, {audit_name} denetimi kapsamında yapılan incelemelerin sonuçlarını içermektedir. Denetim {standard} standartına uygun olarak gerçekleştirilmiş ve toplamda {findings_count} bulgu tespit edilmiştir.",
        "conclusion_text": "Bu denetim raporu, {audit_name} denetimi kapsamında tespit edilen bulguları ve önerileri içermektedir. Organizasyonun ilgili standartlara uygunluğunu artırmak için belirtilen önerilerin değerlendirilmesi ve uygulanması önerilmektedir.",
        "appendix_a": "Ek A: Denetim Metadataları",
        "audit_id": "Denetim ID",
        "created_date": "Oluşturulma Tarihi",
        "last_update": "Son Güncelleme",
        "toc_items": [
            "1. Özet Yönetici Raporu",
            "2. Denetim Kapsamı ve Metodolojisi",
            "3. Bulgular",
            "4. Sonuç ve Öneriler",
            "5. Ekler"
        ],
        "severity_critical": "Kritik",
        "severity_high": "Yüksek",
        "severity_medium": "Orta",
        "severity_low": "Düşük",
        "severity_info": "Bilgi",
        "status_open": "Açık",
        "status_in_progress": "Devam Ediyor",
        "status_resolved": "Çözüldü",
        "status_closed": "Kapatıldı"
    },
    "en": {
        "audit_report": "AUDIT REPORT",
        "executive_summary": "EXECUTIVE SUMMARY",
        "table_of_contents": "TABLE OF CONTENTS",
        "scope_methodology": "AUDIT SCOPE AND METHODOLOGY",
        "findings": "FINDINGS",
        "conclusion": "CONCLUSION AND RECOMMENDATIONS",
        "appendices": "APPENDICES",
        "standard": "Standard",
        "organization": "Organization",
        "project": "Project",
        "audit_date": "Audit Date",
        "report_date": "Report Date",
        "not_specified": "Not Specified",
        "finding_distribution": "Finding Distribution:",
        "finding": "finding(s)",
        "no_findings": "No findings were identified during the audit.",
        "finding_title": "FINDING",
        "severity": "Severity",
        "status_label": "Status",
        "control_reference": "Control Reference",
        "description": "Description",
        "recommendation": "Recommendation",
        "evidence_count": "Evidence Count",
        "evidence_desc": "Description",
        "scope": "Audit Scope",
        "scope_text": "This audit was conducted in accordance with the {standard} standard.",
        "summary_text": "This report contains the results of the {audit_name} audit. The audit was conducted in accordance with the {standard} standard and a total of {findings_count} findings were identified.",
        "conclusion_text": "This audit report contains the findings and recommendations identified during the {audit_name} audit. It is recommended that the organization evaluate and implement the recommendations to improve compliance with relevant standards.",
        "appendix_a": "Appendix A: Audit Metadata",
        "audit_id": "Audit ID",
        "created_date": "Created Date",
        "last_update": "Last Update",
        "toc_items": [
            "1. Executive Summary",
            "2. Audit Scope and Methodology",
            "3. Findings",
            "4. Conclusion and Recommendations",
            "5. Appendices"
        ],
        "severity_critical": "Critical",
        "severity_high": "High",
        "severity_medium": "Medium",
        "severity_low": "Low",
        "severity_info": "Info",
        "status_open": "Open",
        "status_in_progress": "In Progress",
        "status_resolved": "Resolved",
        "status_closed": "Closed"
    }
}

def get_translation(key: str, lang: str = "tr") -> str:
    """Get translation for a key in the specified language"""
    translations = TRANSLATIONS.get(lang, TRANSLATIONS["tr"])
    return translations.get(key, TRANSLATIONS["tr"].get(key, key))

# Severity colors (RGB)
SEVERITY_COLORS = {
    Severity.CRITICAL: RGBColor(220, 53, 69),  # Red
    Severity.HIGH: RGBColor(253, 126, 20),     # Orange
    Severity.MEDIUM: RGBColor(255, 193, 7),    # Yellow
    Severity.LOW: RGBColor(13, 202, 240),      # Cyan
    Severity.INFO: RGBColor(108, 117, 125),    # Gray
}

# Status colors
STATUS_COLORS = {
    Status.OPEN: RGBColor(220, 53, 69),
    Status.IN_PROGRESS: RGBColor(255, 193, 7),
    Status.RESOLVED: RGBColor(22, 197, 94),
    Status.CLOSED: RGBColor(108, 117, 125),
}

def get_severity_text(severity: Severity, lang: str = "tr") -> str:
    severity_keys = {
        Severity.CRITICAL: "severity_critical",
        Severity.HIGH: "severity_high",
        Severity.MEDIUM: "severity_medium",
        Severity.LOW: "severity_low",
        Severity.INFO: "severity_info"
    }
    return get_translation(severity_keys.get(severity, "severity_info"), lang)

def get_status_text(status: Status, lang: str = "tr") -> str:
    status_keys = {
        Status.OPEN: "status_open",
        Status.IN_PROGRESS: "status_in_progress",
        Status.RESOLVED: "status_resolved",
        Status.CLOSED: "status_closed"
    }
    return get_translation(status_keys.get(status, "status_open"), lang)

def add_cover_page(doc: Document, audit: Audit, organization, project, lang: str = "tr"):
    """Add professional cover page"""
    # Cover page section
    section = doc.sections[0]
    section.page_height = Cm(29.7)  # A4 height
    section.page_width = Cm(21.0)   # A4 width
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    
    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(get_translation("audit_report", lang))
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(37, 99, 235)  # Blue
    title_para.space_after = Pt(24)
    
    # Add spacing
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Audit name
    audit_name_para = doc.add_paragraph()
    audit_name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    audit_name_run = audit_name_para.add_run(audit.name)
    audit_name_run.font.size = Pt(20)
    audit_name_run.font.bold = True
    audit_name_para.space_after = Pt(12)
    
    # Standard
    if audit.standard:
        standard_para = doc.add_paragraph()
        standard_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        standard_run = standard_para.add_run(f"{get_translation('standard', lang)}: {audit.standard.value}")
        standard_run.font.size = Pt(14)
        standard_run.font.color.rgb = RGBColor(107, 114, 128)
    
    # Add vertical spacing
    for _ in range(8):
        doc.add_paragraph()
    
    # Information table
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Table content
    date_format = "%d.%m.%Y" if lang == "tr" else "%Y-%m-%d"
    info_data = [
        (get_translation("organization", lang), organization.name),
        (get_translation("project", lang), project.name),
        (get_translation("audit_date", lang), audit.audit_date.strftime(date_format) if audit.audit_date else get_translation("not_specified", lang)),
        (get_translation("report_date", lang), datetime.now().strftime(date_format)),
        (get_translation("standard", lang), audit.standard.value),
    ]
    
    for i, (label, value) in enumerate(info_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = str(value)
        
        # Style label cell
        label_cell = table.rows[i].cells[0]
        label_para = label_cell.paragraphs[0]
        label_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        label_run = label_para.runs[0]
        label_run.font.bold = True
        label_run.font.size = Pt(11)
        
        # Style value cell
        value_cell = table.rows[i].cells[1]
        value_para = value_cell.paragraphs[0]
        value_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        value_run = value_para.runs[0]
        value_run.font.size = Pt(11)
    
    # Add page break
    doc.add_page_break()

def add_executive_summary(doc: Document, audit: Audit, findings, lang: str = "tr"):
    """Add executive summary section"""
    heading = doc.add_heading(get_translation("executive_summary", lang), level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Summary paragraph
    summary_para = doc.add_paragraph()
    summary_text = get_translation("summary_text", lang).format(
        audit_name=audit.name,
        standard=audit.standard.value,
        findings_count=len(findings)
    )
    summary_para.add_run(summary_text).font.size = Pt(11)
    
    # Statistics
    severity_counts = Counter([f.severity for f in findings])
    
    doc.add_paragraph()
    stats_para = doc.add_paragraph(get_translation("finding_distribution", lang), style='List Bullet')
    stats_para.runs[0].font.bold = True
    stats_para.runs[0].font.size = Pt(11)
    
    for severity in [Severity.CRITICAL, Severity.HIGH, Severity.MEDIUM, Severity.LOW, Severity.INFO]:
        count = severity_counts.get(severity, 0)
        if count > 0:
            stat_para = doc.add_paragraph(f"  • {get_severity_text(severity, lang)}: {count} {get_translation('finding', lang)}", style='List Bullet 2')
            stat_para.runs[0].font.size = Pt(10)
    
    doc.add_paragraph()

def add_table_of_contents(doc: Document, lang: str = "tr"):
    """Add table of contents (placeholder)"""
    heading = doc.add_heading(get_translation("table_of_contents", lang), level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    toc_items = TRANSLATIONS.get(lang, TRANSLATIONS["tr"])["toc_items"]
    
    for item in toc_items:
        toc_para = doc.add_paragraph(item)
        toc_para.runs[0].font.size = Pt(11)
        toc_para.space_after = Pt(6)
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_page_break()

def add_methodology_section(doc: Document, audit: Audit, lang: str = "tr"):
    """Add methodology and scope section"""
    heading = doc.add_heading(get_translation("scope_methodology", lang), level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Scope
    scope_para = doc.add_paragraph()
    scope_para.add_run(f"{get_translation('scope', lang)}:").font.bold = True
    scope_para.add_run().font.size = Pt(11)
    
    scope_text = get_translation("scope_text", lang).format(standard=audit.standard.value)
    scope_content = doc.add_paragraph(scope_text)
    scope_content.runs[0].font.size = Pt(11)
    
    if audit.description:
        doc.add_paragraph()
        desc_para = doc.add_paragraph()
        desc_para.add_run(f"{get_translation('description', lang)}:").font.bold = True
        desc_para.add_run().font.size = Pt(11)
        
        desc_content = doc.add_paragraph(audit.description)
        desc_content.runs[0].font.size = Pt(11)
    
    doc.add_paragraph()

def add_findings_section(doc: Document, findings, lang: str = "tr"):
    """Add detailed findings section"""
    heading = doc.add_heading(get_translation("findings", lang), level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    if not findings:
        no_findings = doc.add_paragraph(get_translation("no_findings", lang))
        no_findings.runs[0].font.size = Pt(11)
        no_findings.runs[0].italic = True
        return
    
    for idx, finding in enumerate(findings, 1):
        # Finding heading
        finding_heading = doc.add_heading(f"{get_translation('finding_title', lang)} #{idx}: {finding.title}", level=2)
        finding_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Severity and Status badges
        badge_para = doc.add_paragraph()
        badge_para.paragraph_format.space_after = Pt(6)
        
        # Severity badge
        severity_run = badge_para.add_run(f"{get_translation('severity', lang)}: {get_severity_text(finding.severity, lang)}")
        severity_run.font.bold = True
        severity_run.font.size = Pt(10)
        severity_run.font.color.rgb = SEVERITY_COLORS.get(finding.severity, RGBColor(0, 0, 0))
        severity_run.add_text("  |  ")
        
        # Status badge
        status_run = badge_para.add_run(f"{get_translation('status_label', lang)}: {get_status_text(finding.status, lang)}")
        status_run.font.bold = True
        status_run.font.size = Pt(10)
        status_run.font.color.rgb = STATUS_COLORS.get(finding.status, RGBColor(0, 0, 0))
        
        # Control reference
        if finding.control_reference:
            ref_para = doc.add_paragraph()
            ref_para.add_run(f"{get_translation('control_reference', lang)}: ").font.bold = True
            ref_para.add_run(finding.control_reference).font.size = Pt(11)
        
        # Description
        if finding.description:
            desc_para = doc.add_paragraph()
            desc_para.add_run(f"{get_translation('description', lang)}:").font.bold = True
            desc_para.add_run().font.size = Pt(11)
            
            desc_content = doc.add_paragraph(finding.description)
            desc_content.runs[0].font.size = Pt(11)
            desc_content.paragraph_format.space_after = Pt(6)
        
        # Recommendation
        if finding.recommendation:
            rec_para = doc.add_paragraph()
            rec_para.add_run(f"{get_translation('recommendation', lang)}:").font.bold = True
            rec_para.add_run().font.size = Pt(11)
            
            rec_content = doc.add_paragraph(finding.recommendation)
            rec_content.runs[0].font.size = Pt(11)
            rec_content.paragraph_format.space_after = Pt(6)
        
        # Evidence count
        if finding.evidences:
            evid_para = doc.add_paragraph()
            evid_para.add_run(f"{get_translation('evidence_count', lang)}: {len(finding.evidences)}").font.bold = True
            evid_para.add_run().font.size = Pt(11)
            
            for evidence in finding.evidences:
                # Get file extension
                file_ext = os.path.splitext(evidence.file_name)[1].lower()
                
                # Check if it's an image that can be embedded
                if file_ext in IMAGE_EXTENSIONS:
                    # Try to embed the image
                    file_path = os.path.join(settings.UPLOAD_DIR, evidence.file_path)
                    if os.path.exists(file_path):
                        try:
                            # Add image with max width of 5 inches
                            doc.add_paragraph()  # Add spacing before image
                            img_para = doc.add_paragraph()
                            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = img_para.add_run()
                            run.add_picture(file_path, width=Inches(5))
                            
                            # Add caption below image
                            caption_para = doc.add_paragraph()
                            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            caption_run = caption_para.add_run(f"📷 {evidence.file_name}")
                            caption_run.font.size = Pt(9)
                            caption_run.font.italic = True
                            caption_run.font.color.rgb = RGBColor(107, 114, 128)
                            
                            if evidence.description:
                                desc_para = doc.add_paragraph()
                                desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                desc_run = desc_para.add_run(evidence.description)
                                desc_run.font.size = Pt(9)
                                desc_run.font.color.rgb = RGBColor(107, 114, 128)
                        except Exception as e:
                            # If image embedding fails, fall back to text listing
                            evid_item = doc.add_paragraph(f"  • {evidence.file_name} (görüntü yüklenemedi)", style='List Bullet')
                            evid_item.runs[0].font.size = Pt(10)
                    else:
                        # File not found, list as text
                        evid_item = doc.add_paragraph(f"  • {evidence.file_name} (dosya bulunamadı)", style='List Bullet')
                        evid_item.runs[0].font.size = Pt(10)
                else:
                    # Non-image files: list as before
                    evid_item = doc.add_paragraph(f"  • {evidence.file_name}", style='List Bullet')
                    evid_item.runs[0].font.size = Pt(10)
                    if evidence.description:
                        evid_desc = doc.add_paragraph(f"    {get_translation('evidence_desc', lang)}: {evidence.description}", style='List Bullet 2')
                        evid_desc.runs[0].font.size = Pt(9)
                        evid_desc.runs[0].font.color.rgb = RGBColor(107, 114, 128)
        
        # Add spacing between findings
        doc.add_paragraph()
        doc.add_paragraph()

def add_conclusion_section(doc: Document, audit: Audit, findings, lang: str = "tr"):
    """Add conclusion and recommendations section"""
    heading = doc.add_heading(get_translation("conclusion", lang), level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    conclusion_para = doc.add_paragraph()
    conclusion_text = get_translation("conclusion_text", lang).format(audit_name=audit.name)
    conclusion_para.add_run(conclusion_text).font.size = Pt(11)
    
    doc.add_paragraph()

def add_appendix_section(doc: Document, audit: Audit, lang: str = "tr"):
    """Add appendix section"""
    heading = doc.add_heading(get_translation("appendices", lang), level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    appendix_para = doc.add_paragraph()
    appendix_para.add_run(get_translation("appendix_a", lang)).font.bold = True
    appendix_para.add_run().font.size = Pt(11)
    
    date_format = "%d.%m.%Y %H:%M" if lang == "tr" else "%Y-%m-%d %H:%M"
    metadata_items = [
        (get_translation("audit_id", lang), str(audit.id)),
        (get_translation("created_date", lang), audit.created_at.strftime(date_format) if audit.created_at else "N/A"),
        (get_translation("last_update", lang), audit.updated_at.strftime(date_format) if audit.updated_at else "N/A"),
    ]
    
    for label, value in metadata_items:
        meta_para = doc.add_paragraph(f"  {label}: {value}")
        meta_para.runs[0].font.size = Pt(10)

def generate_audit_word_report(audit: Audit, db: Session) -> str:
    """
    Generate professional Word (.docx) report for an audit
    Follows international audit report standards with cover page
    """
    
    # Get audit details with relationships
    findings = audit.findings
    project = audit.project
    organization = project.organization
    
    # Get audit language (default to 'tr' if not set)
    lang = getattr(audit, 'language', 'tr') or 'tr'
    if lang not in ['tr', 'en']:
        lang = 'tr'
    
    # Create document
    doc = Document()
    
    # Set document properties
    report_title = get_translation("audit_report", lang)
    doc.core_properties.title = f"{report_title} - {audit.name}"
    doc.core_properties.author = "ArchRampart Audit Tool"
    doc.core_properties.comments = f"{audit.standard.value} {report_title}"
    
    # Add cover page
    add_cover_page(doc, audit, organization, project, lang)
    
    # Add table of contents
    add_table_of_contents(doc, lang)
    
    # Add executive summary
    add_executive_summary(doc, audit, findings, lang)
    
    # Add methodology
    add_methodology_section(doc, audit, lang)
    
    # Add findings
    add_findings_section(doc, findings, lang)
    
    # Add conclusion
    add_conclusion_section(doc, audit, findings, lang)
    
    # Add appendix
    add_appendix_section(doc, audit, lang)
    
    # Save to temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    temp_path = temp_file.name
    temp_file.close()
    
    doc.save(temp_path)
    
    return temp_path

